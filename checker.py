"""
STARLUX EMD SOP Format Checker
Based on SOP-EMD-22-101 Rev.05
"""

import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

# ─── Constants ────────────────────────────────────────────────────────────────

DEPT_CODES = {
    "EMD": "機務處",
    "EQP": "品質部",
    "EEP": "工程部",
    "ECP": "修管部",
    "EBP": "客艙維修部",
    "EOP": "組件維修部",
    "ELP": "補給部",
    "ESP": "基地維修部",
    "EMP": "機坪維修部",
}

TIER2_NAMES = ["程序", "辦法", "規範", "Process", "Management", "Management Rule", "Rule"]
TIER3_NAMES = ["程序", "指導書", "作業手冊", "作業辦法", "Procedure", "Instruction",
               "Operation Guide", "Operation Management", "Guide"]

REQUIRED_SECTIONS = [
    ("Revision Summary", "修訂紀要"),
    ("Objectives", "目的"),
    ("Applicability", "範圍"),
    ("Responsibility", "權責"),
    ("Definition", "定義"),
    ("Procedures", "作業內容"),
    ("References", "參考文件"),
    ("Records", "紀錄"),
    ("Attachments", "附件"),
]

COVER_FIELDS = [
    ("文件編號", "SOP No"),
    ("文件名稱", None),
    ("Subject", None),
    ("Classification Level", "機密等級"),
    ("Effective Date", "生效日期"),
    ("Rev. No", "版次"),
    ("Prepared By", "編寫者"),
    ("Checked By", "確認者"),
    ("Approved By", "核可主管"),
]

DATE_PATTERN = re.compile(
    r'\b(\d{4}[/-]\d{1,2}[/-]\d{1,2}|\d{1,2}[/-]\d{1,2}[/-]\d{4}|'
    r'\d{1,2}/\d{1,2}/\d{4}|\d{4}\.\d{1,2}\.\d{1,2})\b'
)
VALID_DATE_PATTERN = re.compile(
    r'\b(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)/\d{2}/\d{4}\b'
)
SOP_NO_PATTERN = re.compile(r'SOP-([A-Z]{3})-(\d{2})-(\d{3})')

MONTHS = ["JAN","FEB","MAR","APR","MAY","JUN","JUL","AUG","SEP","OCT","NOV","DEC"]


# ─── Helpers ──────────────────────────────────────────────────────────────────

def para_text(para):
    return para.text.strip()

def all_paragraphs(doc):
    """Return all paragraphs including inside tables."""
    paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    return paras

def get_run_font(run):
    font = run.font
    name = font.name or ""
    size = font.size
    size_pt = size.pt if size else None
    return name, size_pt

def find_cover_text(doc):
    """Extract first ~20 paragraphs as cover page text."""
    texts = []
    for i, para in enumerate(doc.paragraphs[:25]):
        texts.append(para.text.strip())
    return "\n".join(texts)


# ─── Main Checker ─────────────────────────────────────────────────────────────

def check_sop(docx_path):
    doc = Document(docx_path)
    issues = []  # list of dicts: {id, category, severity, location, description, auto_fixable, fix_type}

    issue_id = [0]
    def add_issue(category, severity, location, description, auto_fixable=False, fix_type=None, fix_data=None):
        issue_id[0] += 1
        issues.append({
            "id": issue_id[0],
            "category": category,
            "severity": severity,  # "error" | "warning" | "info"
            "location": location,
            "description": description,
            "auto_fixable": auto_fixable,
            "fix_type": fix_type,
            "fix_data": fix_data or {},
            "user_action": None,  # "fix" | "skip" | "manual"
        })

    cover_text = find_cover_text(doc)
    all_paras = list(doc.paragraphs)

    # ── 1. Document Number ─────────────────────────────────────────────────
    sop_no = None
    dept_code = None
    year_code = None
    doc_revision = None

    for para in all_paras[:25]:
        t = para.text.strip()
        m = SOP_NO_PATTERN.search(t)
        if m:
            sop_no = m.group(0)
            dept_code = m.group(1)
            year_code = m.group(2)
            break

    if not sop_no:
        add_issue("文件編號", "error", "封面頁",
                  "找不到文件編號，或格式不符合 SOP-XXX-YY-SSS（例：SOP-EQP-22-001）")
    else:
        if dept_code not in DEPT_CODES:
            add_issue("文件編號", "error", f"封面頁 > 文件編號 {sop_no}",
                      f"部門代碼「{dept_code}」不在機務處代碼清單中。\n"
                      f"有效代碼：{', '.join(DEPT_CODES.keys())}")

    # ── 2. Revision Number & Year Check ───────────────────────────────────
    rev_no = None
    for para in all_paras[:25]:
        t = para.text.strip()
        m = re.search(r'Rev(?:ision)?[./\s]*No[.:]?\s*[：:]?\s*(\d+)', t, re.IGNORECASE)
        if m:
            rev_no = m.group(1).zfill(2)
            doc_revision = rev_no
            break

    if rev_no is not None and rev_no == "00" and year_code and sop_no:
        # Find effective date year
        eff_year = None
        for para in all_paras[:25]:
            t = para.text
            dm = re.search(r'(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)/\d{2}/(\d{4})', t, re.IGNORECASE)
            if dm:
                eff_year = dm.group(2)[-2:]
                break
        if eff_year and eff_year != year_code:
            add_issue("文件編號", "warning", f"封面頁 > 文件編號",
                      f"文件版次為 00，編號中的年份代碼「{year_code}」與生效日期年份「{eff_year}」不一致，請確認。")

    # ── 3. Document Name / Tier ────────────────────────────────────────────
    doc_name_cn = ""
    doc_name_en = ""
    for para in all_paras[:25]:
        t = para.text.strip()
        if t.startswith("文件名稱") or "文件名稱：" in t or "文件名稱:" in t:
            doc_name_cn = re.sub(r'^文件名稱[：:]\s*', '', t).strip()
        if t.startswith("Subject") or "Subject：" in t or "Subject:" in t:
            doc_name_en = re.sub(r'^Subject[：:]\s*', '', t, flags=re.IGNORECASE).strip()

    if dept_code and (doc_name_cn or doc_name_en):
        combined = doc_name_cn + " " + doc_name_en
        if dept_code == "EMD":
            if not any(w in combined for w in TIER2_NAMES):
                add_issue("文件稱謂", "error", "封面頁 > 文件名稱",
                          f"部門代碼為 EMD，應使用第二階文件稱謂（{' / '.join(TIER2_NAMES[:6])}），"
                          f"但文件名稱「{doc_name_cn}」中未找到符合的稱謂。")
        else:
            if not any(w in combined for w in TIER3_NAMES):
                add_issue("文件稱謂", "error", "封面頁 > 文件名稱",
                          f"部門代碼為 {dept_code}，應使用第三階文件稱謂（{' / '.join(TIER3_NAMES[:6])}），"
                          f"但文件名稱「{doc_name_cn}」中未找到符合的稱謂。")

    # ── 4. Cover Page Required Fields ─────────────────────────────────────
    missing_fields = []
    for field_en, field_cn in COVER_FIELDS:
        found = False
        search_terms = [field_en]
        if field_cn:
            search_terms.append(field_cn)
        for para in all_paras[:30]:
            if any(term.lower() in para.text.lower() for term in search_terms):
                found = True
                break
        if not found:
            label = f"{field_cn}/{field_en}" if field_cn else field_en
            missing_fields.append(label)

    if missing_fields:
        add_issue("封面頁欄位", "error", "封面頁",
                  f"封面頁缺少以下必要欄位：\n" + "\n".join(f"  • {f}" for f in missing_fields))

    # ── 5. Header / Footer Fields ─────────────────────────────────────────
    has_header_content = False
    has_footer_content = False

    for section in doc.sections:
        h = section.header
        if h and any(para.text.strip() for para in h.paragraphs):
            has_header_content = True
        f = section.footer
        if f and any(para.text.strip() for para in f.paragraphs):
            has_footer_content = True

    if not has_header_content:
        add_issue("頁首", "error", "全文頁首",
                  "未偵測到頁首內容。頁首應包含：公司 Logo、文件編號、英文文件名稱、中文文件名稱。")
    else:
        # Check header fields
        header_text = ""
        for section in doc.sections:
            h = section.header
            if h:
                header_text += " ".join(p.text for p in h.paragraphs)
        missing_header = []
        if not (doc_name_en and doc_name_en[:10] in header_text):
            missing_header.append("英文文件名稱")
        if not (doc_name_cn and doc_name_cn[:6] in header_text):
            missing_header.append("中文文件名稱")
        if sop_no and sop_no not in header_text:
            missing_header.append("文件編號")
        if missing_header:
            add_issue("頁首", "warning", "全文頁首",
                      f"頁首可能缺少以下欄位（需人工確認）：{', '.join(missing_header)}")

    if not has_footer_content:
        add_issue("頁尾", "error", "全文頁尾",
                  "未偵測到頁尾內容。頁尾應包含：負責部門、頁碼/總頁數、版次、生效日期。")
    else:
        footer_text = ""
        for section in doc.sections:
            f = section.footer
            if f:
                footer_text += " ".join(p.text for p in f.paragraphs)
        missing_footer = []
        checks = [("Dept", "負責部門"), ("Page", "頁次"), ("Revision", "版次"), ("Effective Date", "生效日期")]
        for en, cn in checks:
            if en.lower() not in footer_text.lower() and cn not in footer_text:
                missing_footer.append(f"{cn}/{en}")
        if missing_footer:
            add_issue("頁尾", "warning", "全文頁尾",
                      f"頁尾可能缺少以下欄位（需人工確認）：{', '.join(missing_footer)}")

    # ── 6. Page Margins ───────────────────────────────────────────────────
    EXPECTED = {"top": 3.75, "bottom": 2.0, "left": 2.0, "right": 2.0}
    for i, section in enumerate(doc.sections):
        margins = {
            "top": round(section.top_margin.cm, 2) if section.top_margin else None,
            "bottom": round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
            "left": round(section.left_margin.cm, 2) if section.left_margin else None,
            "right": round(section.right_margin.cm, 2) if section.right_margin else None,
        }
        wrong = []
        for key, exp in EXPECTED.items():
            actual = margins[key]
            if actual is not None and abs(actual - exp) > 0.1:
                wrong.append(f"{key}：應為 {exp}cm，實際 {actual}cm")
        if wrong:
            add_issue("頁邊距", "warning", f"版面設定（第 {i+1} 節）",
                      "頁邊距不符合規範，請確認後修改：\n" + "\n".join(f"  • {w}" for w in wrong),
                      auto_fixable=False)

    # ── 7. Font Checks ────────────────────────────────────────────────────
    font_issues = []
    cjk_re = re.compile(r'[\u4e00-\u9fff\u3400-\u4dbf]')

    for i, para in enumerate(all_paras):
        # Skip header/footer (they're in sections)
        if not para.text.strip():
            continue
        for j, run in enumerate(para.runs):
            if not run.text.strip():
                continue
            fname, fsize = get_run_font(run)
            text_sample = run.text[:20].replace('\n', ' ')
            has_cjk = bool(cjk_re.search(run.text))

            # Check size (skip very small runs that might be footnotes)
            if fsize and fsize < 10:
                continue  # likely a special element

            if has_cjk:
                if fname and "JhengHei" not in fname and "正黑" not in fname and fname != "":
                    font_issues.append({
                        "para_idx": i,
                        "run_idx": j,
                        "issue": f"中文字體應為「微軟正黑體」，實際為「{fname}」",
                        "text_sample": text_sample,
                        "fix_type": "font_cn",
                    })
                if fsize and abs(fsize - 12) > 0.5:
                    font_issues.append({
                        "para_idx": i,
                        "run_idx": j,
                        "issue": f"字號應為 12pt，實際為 {fsize:.1f}pt",
                        "text_sample": text_sample,
                        "fix_type": "fontsize",
                    })
            else:
                if fname and fname not in ("Arial", "") and "Arial" not in fname:
                    font_issues.append({
                        "para_idx": i,
                        "run_idx": j,
                        "issue": f"英文字體應為「Arial」，實際為「{fname}」",
                        "text_sample": text_sample,
                        "fix_type": "font_en",
                    })
                if fsize and abs(fsize - 12) > 0.5:
                    font_issues.append({
                        "para_idx": i,
                        "run_idx": j,
                        "issue": f"字號應為 12pt，實際為 {fsize:.1f}pt",
                        "text_sample": text_sample,
                        "fix_type": "fontsize",
                    })

    if font_issues:
        # Group by type
        by_type = {}
        for fi in font_issues[:20]:  # cap at 20 to avoid flooding
            key = fi["fix_type"]
            by_type.setdefault(key, []).append(fi)

        for ftype, items in by_type.items():
            label = {"font_cn": "中文字體", "font_en": "英文字體", "fontsize": "字號"}.get(ftype, ftype)
            desc = f"偵測到 {len(items)} 處{label}不符合規範：\n"
            for item in items[:5]:
                desc += f"  • …{item['text_sample']}… → {item['issue']}\n"
            if len(items) > 5:
                desc += f"  • （另有 {len(items)-5} 處，完整修改請使用「全部修改」）\n"
            add_issue("字體/字號", "warning", "全文內容",
                      desc.strip(), auto_fixable=True,
                      fix_type=ftype,
                      fix_data={"items": items})

    # ── 8. Date Format ────────────────────────────────────────────────────
    date_issues = []
    for i, para in enumerate(all_paras):
        t = para.text
        # Find wrong date patterns
        for m in DATE_PATTERN.finditer(t):
            d = m.group()
            if not VALID_DATE_PATTERN.search(d):
                # Try to convert
                converted = try_convert_date(d)
                date_issues.append({
                    "para_idx": i,
                    "original": d,
                    "suggestion": converted,
                    "context": t[:60],
                })

    for di in date_issues[:10]:
        sug = f" → 建議修改為：{di['suggestion']}" if di['suggestion'] else ""
        add_issue("日期格式", "error",
                  f"段落：「{di['context'][:40]}…」",
                  f"日期「{di['original']}」格式不符合規範（應為 MMM/DD/YYYY）{sug}",
                  auto_fixable=bool(di['suggestion']),
                  fix_type="date",
                  fix_data={"original": di['original'], "replacement": di['suggestion'], "para_idx": di['para_idx']})

    # ── 9. Required Sections ──────────────────────────────────────────────
    full_text = "\n".join(p.text for p in all_paras)
    missing_sections = []
    for en, cn in REQUIRED_SECTIONS:
        if en.lower() not in full_text.lower() and cn not in full_text:
            missing_sections.append(f"{en} {cn}")

    if missing_sections:
        add_issue("必要章節", "error", "文件結構",
                  f"文件缺少以下必要章節：\n" + "\n".join(f"  • {s}" for s in missing_sections))

    # ── 10. Revision Summary Table ────────────────────────────────────────
    if doc_revision is not None:
        rev_rows = count_revision_rows(doc)
        rev_int = int(doc_revision)
        if rev_int == 0 and rev_rows == 0:
            add_issue("修訂紀要", "error", "修訂紀要頁",
                      "未在修訂紀要表格中找到任何版本紀錄，請確認修訂紀要是否正確填寫。")
        elif rev_int >= 1 and rev_rows < 2:
            add_issue("修訂紀要", "error", "修訂紀要頁",
                      f"文件版次為 {doc_revision}（非初始版本），修訂紀要應至少保留最新兩個版本的紀錄，"
                      f"但目前僅偵測到 {rev_rows} 筆紀錄。")

    # ── 11. Bilingual Order (EN before CN) ────────────────────────────────
    bilingual_issues = check_bilingual_order(all_paras)
    for bi in bilingual_issues[:5]:
        add_issue("中英文排列", "warning",
                  f"段落：「{bi['text'][:40]}…」",
                  "偵測到中文敘述可能在英文之前出現。依規範英文應在前、中文在後。\n"
                  "（需人工確認，系統無法自動修改語言排列順序）",
                  auto_fixable=False)

    # ── 12. Column Title Capitalization ───────────────────────────────────
    cap_issues = check_capitalization(all_paras)
    for ci in cap_issues[:5]:
        add_issue("英文大小寫", "warning",
                  f"段落：「{ci['text'][:40]}…」",
                  f"欄位標題/名稱英文每字首字母應大寫（如 Part Number），"
                  f"疑似不符：「{ci['phrase']}」",
                  auto_fixable=False)

    return issues


# ─── Helper Functions ─────────────────────────────────────────────────────────

def try_convert_date(d):
    """Try to convert a date string to MMM/DD/YYYY format."""
    month_map = {
        "01": "JAN", "02": "FEB", "03": "MAR", "04": "APR",
        "05": "MAY", "06": "JUN", "07": "JUL", "08": "AUG",
        "09": "SEP", "10": "OCT", "11": "NOV", "12": "DEC",
        "1": "JAN", "2": "FEB", "3": "MAR", "4": "APR",
        "5": "MAY", "6": "JUN", "7": "JUL", "8": "AUG",
        "9": "SEP",
    }
    # Try YYYY/MM/DD or YYYY-MM-DD
    m = re.match(r'(\d{4})[/-](\d{1,2})[/-](\d{1,2})$', d)
    if m:
        y, mo, day = m.group(1), m.group(2), m.group(3)
        mon = month_map.get(mo)
        if mon:
            return f"{mon}/{day.zfill(2)}/{y}"
    # Try DD/MM/YYYY
    m = re.match(r'(\d{1,2})[/-](\d{1,2})[/-](\d{4})$', d)
    if m:
        day, mo, y = m.group(1), m.group(2), m.group(3)
        mon = month_map.get(mo)
        if mon:
            return f"{mon}/{day.zfill(2)}/{y}"
    return None


def count_revision_rows(doc):
    """Count data rows in revision summary table."""
    count = 0
    for table in doc.tables:
        # Check if this looks like revision table
        flat = " ".join(
            cell.text.strip()
            for row in table.rows
            for cell in row.cells
        ).lower()
        if "修訂" in flat or "revision" in flat:
            for row in table.rows[1:]:  # skip header
                row_text = "".join(c.text.strip() for c in row.cells)
                if row_text and re.search(r'\d{2}', row_text):
                    count += 1
    return count


def check_bilingual_order(paras):
    """Detect paragraphs where Chinese appears before English (heuristic)."""
    cjk_re = re.compile(r'[\u4e00-\u9fff]')
    ascii_re = re.compile(r'[A-Za-z]{3,}')
    issues = []
    for para in paras:
        t = para.text.strip()
        if len(t) < 10:
            continue
        # Find first CJK and first long ASCII word positions
        cjk_match = cjk_re.search(t)
        ascii_match = ascii_re.search(t)
        if cjk_match and ascii_match:
            if cjk_match.start() < ascii_match.start() and cjk_match.start() < 5:
                issues.append({"text": t})
    return issues


def check_capitalization(paras):
    """Check column title capitalization (each word should start with capital)."""
    issues = []
    # Look for patterns like "part number" (lowercase words) in short title-like phrases
    pattern = re.compile(r'\b([a-z][a-z]+(?:\s+[a-z][a-z]+)+)\b')
    skip_words = {"and", "or", "the", "of", "in", "to", "a", "an", "for", "by", "at"}
    for para in paras:
        t = para.text.strip()
        if len(t) > 100 or len(t) < 4:
            continue
        for m in pattern.finditer(t):
            phrase = m.group()
            words = phrase.split()
            if all(w in skip_words for w in words):
                continue
            if any(w[0].islower() and w not in skip_words for w in words):
                issues.append({"text": t, "phrase": phrase})
                break
    return issues


# ─── Auto Fix ─────────────────────────────────────────────────────────────────

def apply_fixes(docx_path, issues_to_fix, output_path):
    """Apply auto-fixable issues to the document."""
    doc = Document(docx_path)
    all_paras = list(doc.paragraphs)
    cjk_re = re.compile(r'[\u4e00-\u9fff\u3400-\u4dbf]')

    for issue in issues_to_fix:
        if not issue.get("auto_fixable"):
            continue
        ft = issue.get("fix_type")
        fd = issue.get("fix_data", {})

        if ft in ("font_cn", "font_en", "fontsize"):
            items = fd.get("items", [])
            for item in items:
                pidx = item.get("para_idx")
                ridx = item.get("run_idx")
                if pidx is not None and pidx < len(all_paras):
                    para = all_paras[pidx]
                    if ridx is not None and ridx < len(para.runs):
                        run = para.runs[ridx]
                        has_cjk = bool(cjk_re.search(run.text))
                        if ft == "font_cn" or (ft == "fontsize" and has_cjk):
                            run.font.name = "Microsoft JhengHei"
                        elif ft == "font_en":
                            run.font.name = "Arial"
                        if ft == "fontsize":
                            run.font.size = Pt(12)

        elif ft == "date":
            orig = fd.get("original")
            repl = fd.get("replacement")
            pidx = fd.get("para_idx")
            if orig and repl and pidx is not None and pidx < len(all_paras):
                para = all_paras[pidx]
                for run in para.runs:
                    if orig in run.text:
                        run.text = run.text.replace(orig, repl)

    doc.save(output_path)
    return output_path
